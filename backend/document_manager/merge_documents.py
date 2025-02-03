from docx import Document
from docx.enum.text import WD_BREAK
import os
from flask import jsonify
from copy import deepcopy

def merge_documents(document_paths, repository_path):
    merged_document = Document()

    # Remove the default empty paragraph if it exists
    if merged_document.paragraphs:
        p = merged_document.paragraphs[0]
        p._element.getparent().remove(p._element)

    for idx, doc_name in enumerate(document_paths):
        doc_path = os.path.join(repository_path, doc_name.lstrip("/"))
        if not os.path.exists(doc_path):
            return jsonify({"error": f"File not found: {doc_name}"}), 400

        doc_to_merge = Document(doc_path)

        # Copy body content
        for element in doc_to_merge.element.body:
            if element.tag.endswith('sectPr'):
                continue
            merged_document.element.body.append(deepcopy(element))  # Use deepcopy to include image references

        # Copy image relationships
        for rel in doc_to_merge.part.rels.values():
            if "image" in rel.reltype:
                if rel.rId not in merged_document.part.rels:
                    merged_document.part.relate_to(
                        rel.target_ref,
                        rel.reltype,
                        rel.target_mode,
                        rId=rel.rId
                    )

        # Add page break if not the last document
        if idx < len(document_paths) - 1:
            last_paragraph = merged_document.paragraphs[-1]
            last_paragraph.add_run().add_break(WD_BREAK.PAGE)

    return merged_document
